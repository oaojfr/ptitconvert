"""
Convertisseur de documents pour PtitConvert
Gère la conversion entre PDF, DOCX et TXT
"""

import PyPDF2
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import os
from pathlib import Path
import io

class DocumentConverter:
    """Convertisseur pour les documents"""
    
    SUPPORTED_INPUT_FORMATS = {'.pdf', '.docx', '.txt'}
    SUPPORTED_OUTPUT_FORMATS = {'pdf', 'docx', 'txt'}
    
    def __init__(self):
        """Initialiser le convertisseur de documents"""
        pass
        
    def convert(self, input_path, output_dir, output_format):
        """
        Convertir un document vers le format spécifié
        
        Args:
            input_path (str): Chemin du fichier d'entrée
            output_dir (str): Répertoire de sortie
            output_format (str): Format de sortie ('pdf', 'docx', 'txt')
            
        Returns:
            bool: True si la conversion a réussi, False sinon
        """
        try:
            input_path = Path(input_path)
            output_format = output_format.lower()
            
            # Vérifier que le format d'entrée est supporté
            if input_path.suffix.lower() not in self.SUPPORTED_INPUT_FORMATS:
                print(f"Format d'entrée non supporté: {input_path.suffix}")
                return False
                
            # Vérifier que le format de sortie est supporté
            if output_format not in self.SUPPORTED_OUTPUT_FORMATS:
                print(f"Format de sortie non supporté: {output_format}")
                return False
                
            # Créer le nom de fichier de sortie
            output_name = f"{input_path.stem}.{output_format}"
            output_path = Path(output_dir) / output_name
            
            # Extraire le texte du document source
            text_content = self._extract_text(input_path)
            if text_content is None:
                return False
                
            # Convertir selon le format de sortie
            if output_format == 'pdf':
                return self._create_pdf(text_content, output_path)
            elif output_format == 'docx':
                return self._create_docx(text_content, output_path)
            elif output_format == 'txt':
                return self._create_txt(text_content, output_path)
            else:
                return False
                
        except Exception as e:
            print(f"Erreur lors de la conversion de document: {e}")
            return False
            
    def _extract_text(self, input_path):
        """
        Extraire le texte d'un document
        
        Args:
            input_path (Path): Chemin du fichier d'entrée
            
        Returns:
            str: Texte extrait ou None en cas d'erreur
        """
        try:
            file_ext = input_path.suffix.lower()
            
            if file_ext == '.pdf':
                return self._extract_pdf_text(input_path)
            elif file_ext == '.docx':
                return self._extract_docx_text(input_path)
            elif file_ext == '.txt':
                return self._extract_txt_text(input_path)
            else:
                return None
                
        except Exception as e:
            print(f"Erreur lors de l'extraction du texte: {e}")
            return None
            
    def _extract_pdf_text(self, pdf_path):
        """Extraire le texte d'un PDF"""
        try:
            text_content = []
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                        
            return '\n\n'.join(text_content)
            
        except Exception as e:
            print(f"Erreur lors de l'extraction du PDF: {e}")
            return None
            
    def _extract_docx_text(self, docx_path):
        """Extraire le texte d'un document Word"""
        try:
            doc = Document(docx_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    
            return '\n\n'.join(text_content)
            
        except Exception as e:
            print(f"Erreur lors de l'extraction du DOCX: {e}")
            return None
            
    def _extract_txt_text(self, txt_path):
        """Extraire le texte d'un fichier texte"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read()
                
        except UnicodeDecodeError:
            # Essayer avec d'autres encodages
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            return None
            
        except Exception as e:
            print(f"Erreur lors de l'extraction du TXT: {e}")
            return None
            
    def _create_pdf(self, text_content, output_path):
        """Créer un PDF à partir du texte"""
        try:
            doc = SimpleDocTemplate(str(output_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Diviser le texte en paragraphes
            paragraphs = text_content.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Créer un paragraphe avec style normal
                    para = Paragraph(para_text.replace('\n', '<br/>'), styles['Normal'])
                    story.append(para)
                    story.append(Spacer(1, 12))
                    
            doc.build(story)
            
            print(f"PDF créé: {output_path}")
            return True
            
        except Exception as e:
            print(f"Erreur lors de la création du PDF: {e}")
            return False
            
    def _create_docx(self, text_content, output_path):
        """Créer un document Word à partir du texte"""
        try:
            doc = Document()
            
            # Diviser le texte en paragraphes
            paragraphs = text_content.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text)
                    
            doc.save(str(output_path))
            
            print(f"DOCX créé: {output_path}")
            return True
            
        except Exception as e:
            print(f"Erreur lors de la création du DOCX: {e}")
            return False
            
    def _create_txt(self, text_content, output_path):
        """Créer un fichier texte"""
        try:
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(text_content)
                
            print(f"TXT créé: {output_path}")
            return True
            
        except Exception as e:
            print(f"Erreur lors de la création du TXT: {e}")
            return False
            
    def get_document_info(self, document_path):
        """
        Obtenir les informations d'un document
        
        Args:
            document_path (str): Chemin du document
            
        Returns:
            dict: Informations sur le document
        """
        try:
            path = Path(document_path)
            file_ext = path.suffix.lower()
            
            info = {
                'file_size': path.stat().st_size,
                'file_format': file_ext[1:].upper()
            }
            
            if file_ext == '.pdf':
                with open(document_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    info['page_count'] = len(pdf_reader.pages)
                    
            elif file_ext == '.docx':
                doc = Document(document_path)
                info['paragraph_count'] = len(doc.paragraphs)
                
            elif file_ext == '.txt':
                with open(document_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    info['character_count'] = len(content)
                    info['line_count'] = len(content.splitlines())
                    
            return info
            
        except Exception as e:
            print(f"Erreur lors de la lecture des informations du document: {e}")
            return None
